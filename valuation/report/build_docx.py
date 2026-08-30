"""
Word Builder — tạo báo cáo định giá .docx theo khuôn 11 phần chuẩn CTCK/quỹ
(SPEC PHẦN B), đồng bộ cấu trúc với template.html. File Word để analyst chỉnh
sửa trực tiếp trước khi phát hành.
"""
import os
import logging
from typing import Any, Dict, List

AI_NOTICE = "⚠ Nháp do AI tạo — cần analyst review trước khi phát hành."
logger = logging.getLogger(__name__)


def build_docx_report(
    data: Dict[str, Any],
    proj_headers: List[str],
    proj_rows: List[Dict[str, Any]],
    charts: Dict[str, str] = None,
    output_path: str = None,
) -> bool:
    """
    Tạo tài liệu Word từ dữ liệu báo cáo 11 phần.

    data: dict render (như template.html) — cần thêm các khối `hist`,
          `assumptions`, `wacc_rows`, `consensus`, `scenarios`, `appendix`,
          `narrative`, `flags`, `market_cap`.
    charts: {"football": path, "heatmap": path, "history": path, "profitability": path}
    """
    if output_path is None:
        return False
    charts = charts or {}
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        doc = Document()

        SLATE = RGBColor(15, 23, 42)
        GRAY = RGBColor(100, 116, 139)
        AMBER = RGBColor(146, 64, 14)

        def add_section(text):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.size = Pt(14)
            r.font.bold = True
            r.font.color.rgb = SLATE
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)

        def add_ai_notice():
            if data.get("narrative", {}).get("ai_generated"):
                p = doc.add_paragraph()
                r = p.add_run(AI_NOTICE)
                r.font.size = Pt(9)
                r.font.bold = True
                r.font.color.rgb = AMBER

        def add_narrative(key):
            add_ai_notice()
            text = data.get("narrative", {}).get(key, "")
            for para in str(text).split("\n"):
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        def add_table(headers: List[str], rows: List[Dict[str, Any]], first_col: str):
            t = doc.add_table(rows=len(rows) + 1, cols=len(headers) + 1)
            t.style = "Table Grid"
            t.cell(0, 0).paragraphs[0].add_run(first_col).font.bold = True
            for ci, h in enumerate(headers):
                t.cell(0, ci + 1).paragraphs[0].add_run(str(h)).font.bold = True
            for ri, r in enumerate(rows):
                t.cell(ri + 1, 0).paragraphs[0].add_run(r["label"])
                for ci, val in enumerate(r["values"]):
                    t.cell(ri + 1, ci + 1).paragraphs[0].add_run(str(val))
            return t

        def add_kv_table(rows: List[Dict[str, str]]):
            t = doc.add_table(rows=len(rows), cols=2)
            t.style = "Table Grid"
            for ri, r in enumerate(rows):
                t.cell(ri, 0).paragraphs[0].add_run(r["label"])
                t.cell(ri, 1).paragraphs[0].add_run(str(r["value"])).font.bold = True
            return t

        def add_chart(key, width=6.0):
            path = charts.get(key)
            if path and os.path.exists(path):
                doc.add_paragraph().paragraph_format.space_before = Pt(6)
                doc.add_picture(path, width=Inches(width))

        # ==== 1. COVER ====
        title_p = doc.add_paragraph()
        tr = title_p.add_run("BÁO CÁO PHÂN TÍCH & ĐỊNH GIÁ DOANH NGHIỆP")
        tr.font.size = Pt(19)
        tr.font.bold = True
        tr.font.color.rgb = SLATE
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        meta_p = doc.add_paragraph()
        mr = meta_p.add_run(
            f"{data['ticker']} · {data.get('name', '')} · Ngành: {data.get('sector', '')} | "
            f"Ngày lập: {data['date']} | Người phân tích: {data['analyst']}"
        )
        mr.font.size = Pt(10)
        mr.font.italic = True
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        rec_table = doc.add_table(rows=1, cols=1)
        rec_cell = rec_table.cell(0, 0)
        rec_cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>'))
        rec_p = rec_cell.paragraphs[0]
        r1 = rec_p.add_run("KHUYẾN NGHỊ ĐẦU TƯ\n")
        r1.font.size = Pt(11); r1.font.bold = True; r1.font.color.rgb = GRAY
        rec_label = data["recommendation"]
        r2 = rec_p.add_run(f"{rec_label}\n")
        r2.font.size = Pt(26); r2.font.bold = True
        if rec_label in ("MUA", "KHẢ QUAN"):
            r2.font.color.rgb = RGBColor(16, 185, 129)
        elif rec_label == "NẮM GIỮ":
            r2.font.color.rgb = RGBColor(245, 158, 11)
        else:
            r2.font.color.rgb = RGBColor(239, 68, 68)
        r3 = rec_p.add_run(
            f"Giá mục tiêu: {data['target_price']} VND (Upside: {data['upside']}%)\n"
            f"Thị giá hiện tại: {data['current_price']} VND · "
            f"Vốn hóa: {data.get('market_cap', 'N/A')} tỷ đồng · "
            f"SLCP: {data['shares']} triệu cp"
        )
        r3.font.size = Pt(11); r3.font.color.rgb = SLATE

        # ==== 2. LUẬN ĐIỂM ĐẦU TƯ ====
        add_section("1. Luận điểm đầu tư")
        add_narrative("thesis")

        # ==== 3. TÓM TẮT ĐỊNH GIÁ ====
        add_section("2. Tóm tắt định giá")
        add_table(
            ["Trọng số", "Giá ngụ ý (VND)"],
            [
                {"label": f"Mô hình Nội tại ({data['intrinsic_method']})",
                 "values": [f"{data['weight_intrinsic']}%", data['intrinsic_price']]},
                {"label": f"Mô hình So sánh ({data['relative_method']})",
                 "values": [f"{data['weight_relative']}%", data['relative_price']]},
                {"label": "Giá mục tiêu pha trộn (Blended)",
                 "values": ["100%", data['target_price']]},
            ],
            "Phương pháp",
        )
        add_chart("football")

        # ==== 4-5. TỔNG QUAN DN + NGÀNH ====
        add_section("3. Tổng quan doanh nghiệp")
        add_narrative("overview")
        add_section("4. Bối cảnh ngành")
        add_narrative("industry")
        add_section("4A. Cổ tức, tăng vốn & quyền cổ đông")
        add_narrative("corporate_actions")

        # ==== 6. TÀI CHÍNH LỊCH SỬ ====
        add_section("5. Phân tích tài chính lịch sử")
        hist = data.get("hist", {})
        if hist:
            add_table(hist["headers"], hist["rows"], "Chỉ tiêu (tỷ đồng)")
        add_chart("history")
        add_chart("profitability")

        # ==== 7. GIẢ ĐỊNH DỰ PHÓNG ====
        add_section("6. Giả định dự phóng")
        ass = data.get("assumptions", {})
        if ass:
            add_table(proj_headers, ass["schedule_rows"], "Giả định")
            add_kv_table(ass["single_rows"])

        # ==== 8. CHI TIẾT ĐỊNH GIÁ ====
        add_section("7. Chi tiết định giá")
        doc.add_paragraph("7.1. Bảng dự phóng tài chính 5 năm (tỷ đồng)").runs[0].font.bold = True
        add_table(proj_headers, proj_rows, "Chỉ tiêu")
        doc.add_paragraph("7.2. Bóc tách chi phí vốn (CAPM / WACC)").runs[0].font.bold = True
        add_kv_table(data.get("wacc_rows", []))
        consensus = data.get("consensus")
        if consensus:
            doc.add_paragraph("7.3. So sánh với định giá các công ty chứng khoán").runs[0].font.bold = True
            dev_note = " ⚠ lệch lớn — cần rà soát giả định" if consensus.get("flag_high") else ""
            kv = [
                {"label": "Giá mục tiêu của mô hình VN100", "value": f"{consensus['our_target']:,.0f} VND"},
                {"label": f"Trung vị consensus ({consensus['n_reports']} báo cáo, 180 ngày)",
                 "value": f"{consensus['consensus_median']:,.0f} VND"},
            ]
            if consensus.get("range_min") and consensus.get("range_max"):
                kv.append({"label": f"Dải giá mục tiêu CTCK ({consensus.get('n_brokers', 0)} CTCK)",
                           "value": f"{consensus['range_min']:,.0f} – {consensus['range_max']:,.0f} VND"})
            if consensus.get("current_price"):
                kv.append({"label": "Thị giá hiện tại",
                           "value": f"{consensus['current_price']:,.0f} VND"})
            kv.append({"label": "Chênh lệch mô hình vs trung vị CTCK",
                       "value": f"{consensus['deviation']:+.1%}{dev_note}"})
            add_kv_table(kv)

            # Bảng chi tiết từng CTCK
            broker_rows = consensus.get("broker_rows") or []
            if broker_rows:
                p = doc.add_paragraph("Chi tiết khuyến nghị từng CTCK (mới nhất/CTCK, 180 ngày):")
                p.runs[0].font.bold = True
                add_table(
                    ["Ngày BC", "Giá mục tiêu (VND)", "Khuyến nghị", "Mô hình so với CTCK"],
                    [
                        {"label": b["broker"],
                         "values": [b["report_date"], f"{b['target_price']:,.0f}",
                                    b["rating"], f"{b['vs_model']:+.1%}"]}
                        for b in broker_rows
                    ],
                    "CTCK",
                )

            # AI tổng hợp điểm chung/riêng/mấu chốt
            synth = consensus.get("synthesis")
            if synth:
                p = doc.add_paragraph("Tổng hợp luận điểm các CTCK ")
                p.runs[0].font.bold = True
                note = p.add_run(f"(AI tổng hợp từ báo cáo CTCK — cần analyst review · nguồn: {synth.get('brokers', '')})")
                note.font.size = Pt(9)
                note.font.color.rgb = RGBColor(0xB4, 0x53, 0x09)

                def _bullet_block(title, items, rgb):
                    if not items:
                        return
                    if isinstance(items, str):
                        items = [items]
                    h = doc.add_paragraph(title)
                    h.runs[0].font.bold = True
                    h.runs[0].font.color.rgb = RGBColor(*rgb)
                    for it in items:
                        doc.add_paragraph(str(it), style="List Bullet")

                _bullet_block("Điểm CHUNG các CTCK đồng thuận:", synth.get("diem_chung"), (0x16, 0x65, 0x34))
                _bullet_block("Điểm RIÊNG / khác biệt giữa các CTCK:", synth.get("diem_rieng"), (0x92, 0x40, 0x0E))
                _bullet_block("Điểm MẤU CHỐT — thị trường có thể bỏ sót:", synth.get("diem_mau_chot"), (0x1E, 0x40, 0xAF))
                if synth.get("doi_chieu_noi_bo"):
                    _bullet_block("Đối chiếu mô hình nội bộ vs đồng thuận:",
                                  [synth["doi_chieu_noi_bo"]], (0x5B, 0x21, 0xB6))

        # ==== 9. ĐỘ NHẠY & KỊCH BẢN ====
        add_section("8. Phân tích độ nhạy & kịch bản")
        scenarios = data.get("scenarios")
        if scenarios and scenarios.get("applicable"):
            add_table(
                ["Giá mục tiêu (VND)", "Upside"],
                [
                    {"label": r["scenario"],
                     "values": [f"{r['target']:,.0f}", f"{r['upside']:+.1%}"]}
                    for r in scenarios["rows"]
                ],
                "Kịch bản",
            )
        elif scenarios:
            doc.add_paragraph(
                "Phương pháp định giá hiện tại (proxy từ giá trị sổ sách) không co giãn "
                "theo giả định tăng trưởng/biên — bảng kịch bản Bull/Bear không áp dụng."
            )
        add_chart("heatmap", width=6.2)

        # ==== 10. RỦI RO ====
        add_section("9. Rủi ro đầu tư")
        add_narrative("risks")
        flags = data.get("flags") or []
        if flags:
            p = doc.add_paragraph()
            r = p.add_run("Cờ QC từ hệ thống: " + ", ".join(flags))
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(153, 27, 27)
        if data.get("notes"):
            doc.add_paragraph(f"Ghi chú của analyst: {data['notes']}")

        # ==== 11. PHỤ LỤC ====
        add_section("10. Phụ lục — Báo cáo tài chính lịch sử")
        appendix = data.get("appendix", {})
        if appendix:
            doc.add_paragraph("Kết quả kinh doanh (tỷ đồng)").runs[0].font.bold = True
            add_table(appendix["headers"], appendix["income_statement"], "Chỉ tiêu")
            doc.add_paragraph("Cân đối kế toán (tỷ đồng)").runs[0].font.bold = True
            add_table(appendix["headers"], appendix["balance_sheet"], "Chỉ tiêu")
            if appendix.get("cash_flow"):
                doc.add_paragraph("Lưu chuyển tiền tệ (tỷ đồng)").runs[0].font.bold = True
                add_table(appendix["headers"], appendix["cash_flow"], "Chỉ tiêu")

        # Footer disclaimer
        doc.add_paragraph().paragraph_format.space_before = Pt(24)
        fp = doc.add_paragraph()
        fr = fp.add_run(
            "Miễn trừ trách nhiệm: báo cáo được tạo tự động bởi Hệ thống định giá VN100, là bản nháp "
            "phục vụ phân tích nội bộ. Văn bản do AI sinh cần được analyst thẩm định. Kết quả phụ thuộc "
            "giả định đầu vào; không cấu thành khuyến nghị mua/bán cho bên thứ ba."
        )
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor(148, 163, 184)
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(output_path)
        logger.info("Word report created at %s", output_path)
        return True
    except Exception as e:
        logger.warning("Không tạo được Word report: %s", e)
        return False
